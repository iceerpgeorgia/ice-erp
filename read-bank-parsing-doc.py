import docx
from pathlib import Path

# Load the document
doc_path = Path(r'C:\next-postgres-starter\Prompts\Bank Accounts Parsing.docx')
doc = docx.Document(doc_path)

print("=" * 80)
print("BANK ACCOUNTS PARSING DOCUMENT ANALYSIS")
print("=" * 80)
print()

# Extract all text content
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(para.text)
        print()

# Check if there are tables
if doc.tables:
    print("\n" + "=" * 80)
    print(f"FOUND {len(doc.tables)} TABLE(S) IN DOCUMENT")
    print("=" * 80)
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            print(" | ".join(row_data))
